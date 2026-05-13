"""
Génération des rapports benchmark de transformation organisationnelle.
Formats : DOCX, HTML (PDF optionnel via weasyprint).
"""
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List

logger = logging.getLogger(__name__)

# Palettes couleurs
IMPORTANCE_COLORS  = {"Critique": "C0392B", "Élevée": "E67E22", "Modérée": "27AE60"}
PRIORITY_COLORS    = {"Haute": "C0392B", "Moyenne": "E67E22", "Faible": "27AE60"}
MATURITE_COLORS    = {"Mature/Répandue": "27AE60", "En développement": "E67E22", "Émergente": "2E86C1"}
ADOPTION_COLORS    = {"Majoritaire (>60%)": "27AE60", "En diffusion (20-60%)": "E67E22", "Pionnier (<20%)": "2E86C1"}
ENGAGEMENT_COLORS  = {"Fort": "27AE60", "Modéré": "E67E22", "Émergent": "2E86C1"}
DIRECTION_COLORS   = {
    "Vers plus d'externalisation": "E67E22",
    "Vers plus d'internalisation": "2E86C1",
    "Nouveaux modèles hybrides": "8E44AD",
    "Stable": "7F8C8D",
}
SCENARIO_COLORS    = {"Optimiste": "27AE60", "Central": "2E86C1", "Pessimiste": "C0392B"}
NIVEAUX_FCS_ORDER  = ["Stratégique", "Organisationnel", "Opérationnel", "Technologique", "Humain & RH"]
TYPE_REC_COLORS    = {
    "Stratégique": "1A5F8A", "Organisationnel": "2E86C1", "RH & Compétences": "8E44AD",
    "Digital": "16A085", "Gouvernance": "D35400", "RSE": "27AE60",
}


# ─────────────────────────────────────────
#  Helpers DOCX
# ─────────────────────────────────────────

def _shd_cell(cell, hex_color: str, text: str = "", bold: bool = True, font_size: int = 9):
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)
    if text:
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = bold
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(font_size)


def _heading(doc, text: str, level: int):
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    color = RGBColor(0x1A, 0x5F, 0x8A) if level <= 2 else RGBColor(0x2E, 0x86, 0xC1)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _table_header(table, headers: List[str], bg: str = "1A5F8A"):
    from docx.shared import RGBColor, Pt
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _shd_cell(cell, bg, h, bold=True, font_size=10)


def _sources_str(ids: List[int]) -> str:
    if not ids:
        return ""
    return " [" + ", ".join(str(i) for i in ids) + "]"


def _safe_list(val) -> list:
    """
    Normalise un champ qui devrait être une liste de dicts.
    Claude retourne parfois une string JSON — on la parse automatiquement.
    """
    import json as _json
    if isinstance(val, str):
        try:
            parsed = _json.loads(val)
            if isinstance(parsed, list):
                return [item for item in parsed if isinstance(item, dict)]
        except Exception:
            pass
        return []
    if isinstance(val, list):
        return [item for item in val if isinstance(item, dict)]
    return []


# ─────────────────────────────────────────
#  DOCX
# ─────────────────────────────────────────

def generate_docx(analysis: Dict[str, Any], output_path: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    meta = analysis.get("_meta", {})
    period = meta.get("period", datetime.now().strftime("%B %Y"))
    sector = meta.get("sector", "Secteur")
    nb_sources = meta.get("nb_sources_analysees", "N/A")

    # ── Page de titre ──
    for text, size, bold, color in [
        ("BENCHMARK DE TRANSFORMATION ORGANISATIONNELLE", 20, True, RGBColor(0x1A, 0x5F, 0x8A)),
        (sector, 15, False, RGBColor(0x55, 0x55, 0x55)),
        (f"Période : {period}", 11, False, RGBColor(0x77, 0x77, 0x77)),
        (f"Sources analysées : {nb_sources} — Généré le {datetime.now().strftime('%d/%m/%Y')}", 10, False, RGBColor(0x99, 0x99, 0x99)),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.color.rgb = color

    # Qualité sources
    qs = analysis.get("qualite_sources", {})
    if qs:
        qp = doc.add_paragraph()
        qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fiab = qs.get("fiabilite_globale", "")
        qp.add_run(f"Fiabilité des sources : {fiab}").font.size = Pt(10)

    doc.add_page_break()

    # ── Synthèse exécutive ──
    _heading(doc, "Synthèse Exécutive", 1)
    se = analysis.get("synthese_executive", {})
    texte = se.get("texte", "") if isinstance(se, dict) else str(se)
    src_ids = se.get("sources", []) if isinstance(se, dict) else []
    doc.add_paragraph(texte + _sources_str(src_ids))

    if qs.get("note_methodologique"):
        p = doc.add_paragraph()
        p.add_run("Note méthodologique : ").bold = True
        p.add_run(qs["note_methodologique"])
    doc.add_paragraph()

    # ── Facteurs Clés de Succès ──
    _heading(doc, "Facteurs Clés de Succès", 1)
    fcs_list = _safe_list(analysis.get("facteurs_cles_succes", []))
    for niveau in NIVEAUX_FCS_ORDER:
        items = [f for f in fcs_list if f.get("niveau") == niveau]
        if not items:
            continue
        _heading(doc, niveau, 2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        _table_header(table, ["Facteur", "Description", "Importance"])
        for item in items:
            row = table.add_row().cells
            row[0].text = item.get("facteur", "") + _sources_str(item.get("sources", []))
            row[1].text = item.get("description", "")
            imp = item.get("importance", "")
            _shd_cell(row[2], IMPORTANCE_COLORS.get(imp, "888888"), imp)
        doc.add_paragraph()

    # ── Dimensionnement ──
    _heading(doc, "Tendances de Dimensionnement", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _table_header(table, ["Tendance", "Description", "Impact effectifs", "Fonctions"])
    for item in _safe_list(analysis.get("tendances_dimensionnement", [])):
        row = table.add_row().cells
        row[0].text = item.get("tendance", "") + _sources_str(item.get("sources", []))
        row[1].text = item.get("description", "")
        impact = item.get("impact_effectifs", "")
        row[2].text = impact
        row[3].text = ", ".join(item.get("fonctions_concernees", []))
    doc.add_paragraph()

    # ── Gouvernance ──
    _heading(doc, "Pratiques de Gouvernance", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _table_header(table, ["Pratique", "Description", "Maturité"])
    for item in _safe_list(analysis.get("pratiques_gouvernance", [])):
        row = table.add_row().cells
        row[0].text = item.get("pratique", "") + _sources_str(item.get("sources", []))
        row[1].text = item.get("description", "")
        mat = item.get("maturite", "")
        _shd_cell(row[2], MATURITE_COLORS.get(mat, "888888"), mat, font_size=8)
    doc.add_paragraph()

    # ── Gestion de la performance ──
    _heading(doc, "Gestion de la Performance", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _table_header(table, ["Pratique", "Description", "Adoption"])
    for item in _safe_list(analysis.get("gestion_performance", [])):
        row = table.add_row().cells
        row[0].text = item.get("pratique", "") + _sources_str(item.get("sources", []))
        row[1].text = item.get("description", "")
        adp = item.get("niveau_adoption", "")
        _shd_cell(row[2], ADOPTION_COLORS.get(adp, "888888"), adp, font_size=8)
    doc.add_paragraph()

    # ── Externalisation ──
    _heading(doc, "Externalisation & Partenariats", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _table_header(table, ["Domaine", "Tendance", "Direction"])
    for item in _safe_list(analysis.get("externalisation_partenariats", [])):
        row = table.add_row().cells
        row[0].text = item.get("domaine", "") + _sources_str(item.get("sources", []))
        row[1].text = item.get("tendance", "")
        dir_ = item.get("direction", "")
        _shd_cell(row[2], DIRECTION_COLORS.get(dir_, "888888"), dir_, font_size=8)
    doc.add_paragraph()

    # ── RSE & Éthique ──
    _heading(doc, "RSE & Éthique", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    _table_header(table, ["Axe", "Description", "Engagement"])
    for item in _safe_list(analysis.get("rse_ethique", [])):
        row = table.add_row().cells
        row[0].text = item.get("axe", "") + _sources_str(item.get("sources", []))
        row[1].text = item.get("description", "")
        eng = item.get("niveau_engagement", "")
        _shd_cell(row[2], ENGAGEMENT_COLORS.get(eng, "888888"), eng)
    doc.add_paragraph()

    # ── Signaux faibles ──
    _heading(doc, "Signaux Faibles & Disruptions Émergentes", 1)
    for sf in _safe_list(analysis.get("signaux_faibles", [])):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(sf.get("signal", "") + _sources_str(sf.get("sources", []))).bold = True
        doc.add_paragraph(
            f"Implication : {sf.get('implication_organisationnelle', '')}  |  "
            f"Horizon : {sf.get('horizon_emergence', '')}"
        )
    doc.add_paragraph()

    # ── Prospective ──
    _heading(doc, "Analyse Prospective", 1)
    prospective = analysis.get("prospective", {})
    for key, label in [("horizon_court_terme", "Court terme"), ("horizon_moyen_terme", "Moyen terme")]:
        h = prospective.get(key, {})
        if h:
            _heading(doc, f"{label} ({h.get('periode', '')})", 2)
            doc.add_paragraph("Évolutions probables :")
            for e in h.get("evolutions_probables", []):
                doc.add_paragraph(e, style="List Bullet")
            doc.add_paragraph("Risques principaux :")
            for r in h.get("risques_principaux", []):
                doc.add_paragraph(r, style="List Bullet")
            doc.add_paragraph()

    scenarios = prospective.get("scenarios", [])
    if scenarios:
        _heading(doc, "Scénarios", 2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        _table_header(table, ["Scénario", "Description", "Conditions", "Implications organisationnelles"])
        for s in scenarios:
            row = table.add_row().cells
            nom = s.get("nom", "")
            _shd_cell(row[0], SCENARIO_COLORS.get(nom, "888888"), nom)
            row[1].text = s.get("description", "")
            row[2].text = s.get("conditions_realisation", "")
            row[3].text = s.get("implications_organisationnelles", "")
        doc.add_paragraph()

    # ── Recommandations ──
    _heading(doc, "Recommandations Stratégiques", 1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    _table_header(table, ["Action", "Justification", "Type", "Priorité", "Horizon"])
    for r in _safe_list(analysis.get("recommandations", [])):
        row = table.add_row().cells
        row[0].text = r.get("action", "") + _sources_str(r.get("sources", []))
        row[1].text = r.get("justification", "")
        type_ = r.get("type", "")
        _shd_cell(row[2], TYPE_REC_COLORS.get(type_, "888888"), type_, font_size=8)
        prio = r.get("priorite", "")
        _shd_cell(row[3], PRIORITY_COLORS.get(prio, "888888"), prio)
        row[4].text = r.get("horizon", "")
    doc.add_paragraph()

    # ── Index des sources ──
    _heading(doc, "Sources & Méthodologie", 1)
    index = analysis.get("index_sources", [])
    if index:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        _table_header(table, ["N°", "Titre", "Source", "Date"])
        for s in sorted(index, key=lambda x: x.get("id", 0)):
            row = table.add_row().cells
            row[0].text = f"[{s.get('id', '')}]"
            title = s.get("titre", "")
            url = s.get("url", "")
            row[1].text = f"{title}\n{url}" if url else title
            row[2].text = s.get("source", "")
            row[3].text = s.get("date", "")

    # ── Footer ──
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.add_run(
            f"Benchmark confidentiel — {sector} — {period} — "
            f"Généré le {datetime.now().strftime('%d/%m/%Y')}"
        ).font.size = Pt(8)

    doc.save(output_path)
    logger.info(f"Rapport DOCX généré : {output_path}")
    return output_path


# ─────────────────────────────────────────
#  HTML
# ─────────────────────────────────────────

def generate_html(analysis: Dict[str, Any], output_path: str, template_path: str) -> str:
    from jinja2 import Environment, FileSystemLoader

    env = Environment(
        loader=FileSystemLoader(str(Path(template_path).parent)),
        autoescape=True,
    )
    template = env.get_template(Path(template_path).name)
    meta = analysis.get("_meta", {})

    html = template.render(
        analysis=analysis,
        sector=meta.get("sector", ""),
        period=meta.get("period", ""),
        generated_at=datetime.now().strftime("%d/%m/%Y à %H:%M"),
        nb_sources=meta.get("nb_sources_analysees", 0),
        freshness=meta.get("freshness", {}),
        niveaux_fcs=NIVEAUX_FCS_ORDER,
        importance_colors=IMPORTANCE_COLORS,
        priority_colors=PRIORITY_COLORS,
        maturite_colors=MATURITE_COLORS,
        adoption_colors=ADOPTION_COLORS,
        engagement_colors=ENGAGEMENT_COLORS,
        direction_colors=DIRECTION_COLORS,
        scenario_colors=SCENARIO_COLORS,
        type_rec_colors=TYPE_REC_COLORS,
    )

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)

    logger.info(f"Rapport HTML généré : {output_path}")
    return output_path


# ─────────────────────────────────────────
#  PDF optionnel
# ─────────────────────────────────────────

def generate_pdf(html_path: str, output_path: str) -> str | None:
    try:
        import weasyprint
        weasyprint.HTML(filename=html_path).write_pdf(output_path)
        logger.info(f"Rapport PDF généré : {output_path}")
        return output_path
    except ImportError:
        logger.warning("weasyprint non installé — PDF ignoré. (pip install weasyprint)")
        return None
    except Exception as e:
        logger.error(f"Erreur génération PDF : {e}")
        return None


# ─────────────────────────────────────────
#  Point d'entrée
# ─────────────────────────────────────────

def generate_reports(analysis: Dict[str, Any], settings: Dict, project_root: str) -> List[str]:
    import json

    formats = settings.get("reporting", {}).get("formats", ["docx", "html"])
    output_dir = Path(project_root) / settings.get("reporting", {}).get("output_dir", "reports")
    output_dir.mkdir(parents=True, exist_ok=True)

    meta = analysis.get("_meta", {})
    sector_slug = (
        meta.get("sector", "secteur")
        .lower()
        .replace(" ", "_")
        .replace("&", "")
        .replace("—", "")
        .replace("–", "")
        .strip("_")
    )
    date_slug = datetime.now().strftime("%Y-%m")
    pattern = settings.get("reporting", {}).get("filename_pattern", "veille_{sector}_{date}")
    base_name = pattern.replace("{sector}", sector_slug).replace("{date}", date_slug)

    template_path = str(Path(project_root) / "templates" / "report.html")
    generated = []

    # ── Sauvegarde JSON (permet de régénérer les rapports sans rappeler Claude) ──
    json_path = output_dir / f"{base_name}.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(analysis, f, ensure_ascii=False, indent=2)
    logger.info(f"Données JSON sauvegardées : {json_path}")

    if "html" in formats or "pdf" in formats:
        html_path = str(output_dir / f"{base_name}.html")
        generate_html(analysis, html_path, template_path)
        generated.append(html_path)
        if "pdf" in formats:
            pdf_path = str(output_dir / f"{base_name}.pdf")
            r = generate_pdf(html_path, pdf_path)
            if r:
                generated.append(r)

    if "docx" in formats:
        docx_path = str(output_dir / f"{base_name}.docx")
        generate_docx(analysis, docx_path)
        generated.append(docx_path)

    return generated
